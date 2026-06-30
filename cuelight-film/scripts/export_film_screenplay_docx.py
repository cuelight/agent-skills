#!/usr/bin/env python3
import argparse
import json
import shutil
import subprocess
import sys
from pathlib import Path

try:
    import yaml
except ImportError:
    yaml = None

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ImportError:
    Document = None


PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
STYLE_NAMES = {
    "scene_heading": "CueLight Scene Heading",
    "action": "CueLight Action",
    "speaker": "CueLight Character",
    "dialogue": "CueLight Dialogue",
    "parenthetical": "CueLight Parenthetical",
    "transition": "CueLight Transition",
    "special": "CueLight Special",
    "note": "CueLight Note",
}


def make_issue(level, issue_type, message, path=None):
    issue = {"level": level, "type": issue_type, "message": message}
    if path:
        issue["path"] = str(path).replace("\\", "/")
    return issue


def read_yaml(path, issues):
    if yaml is None:
        issues.append(make_issue("error", "pyyaml_missing", "PyYAML is required.", path))
        return None
    try:
        return yaml.safe_load(path.read_text(encoding="utf-8"))
    except Exception as exc:
        issues.append(make_issue("error", "yaml_parse_error", f"Cannot parse YAML: {exc}", path))
        return None


def ensure_docx_available():
    if Document is None:
        print(
            "ERROR: python-docx is required. Use the Codex bundled Python or install python-docx.",
            file=sys.stderr,
        )
        return False
    return True


def set_run_font(run, size_pt=11, bold=False):
    run.font.name = "SimSun"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    r_pr = run._element.rPr
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        run._element.insert(0, r_pr)
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Courier New")
    r_fonts.set(qn("w:hAnsi"), "Courier New")
    r_fonts.set(qn("w:eastAsia"), "SimSun")


def set_paragraph_format(paragraph, *, left_cm=0, right_cm=0, first_line_cm=0, before_pt=0, after_pt=6, align=None):
    fmt = paragraph.paragraph_format
    fmt.left_indent = Cm(left_cm)
    fmt.right_indent = Cm(right_cm)
    fmt.first_line_indent = Cm(first_line_cm)
    fmt.space_before = Pt(before_pt)
    fmt.space_after = Pt(after_pt)
    fmt.line_spacing = 1.08
    if align is not None:
        paragraph.alignment = align


def ensure_style(doc, name):
    styles = doc.styles
    if name in styles:
        return styles[name]
    return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def configure_style(doc, name, *, left_cm=0, right_cm=0, first_line_cm=0, before_pt=0, after_pt=6, align=None, size_pt=11, bold=False, italic=False):
    style = ensure_style(doc, name)
    style.font.name = "SimSun"
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.italic = italic
    style._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    fmt = style.paragraph_format
    fmt.left_indent = Cm(left_cm)
    fmt.right_indent = Cm(right_cm)
    fmt.first_line_indent = Cm(first_line_cm)
    fmt.space_before = Pt(before_pt)
    fmt.space_after = Pt(after_pt)
    fmt.line_spacing = 1.08
    if align is not None:
        style.paragraph_format.alignment = align


def configure_screenplay_styles(doc):
    configure_style(doc, STYLE_NAMES["scene_heading"], before_pt=10, after_pt=8, size_pt=12, bold=True)
    configure_style(doc, STYLE_NAMES["action"], after_pt=7, size_pt=11)
    configure_style(doc, STYLE_NAMES["speaker"], left_cm=6.6, after_pt=1, size_pt=11, bold=True)
    configure_style(doc, STYLE_NAMES["dialogue"], left_cm=4.1, right_cm=3.6, after_pt=5, size_pt=11)
    configure_style(doc, STYLE_NAMES["parenthetical"], left_cm=5.0, right_cm=4.0, after_pt=3, size_pt=10)
    configure_style(doc, STYLE_NAMES["transition"], after_pt=8, align=WD_ALIGN_PARAGRAPH.RIGHT, size_pt=11)
    configure_style(doc, STYLE_NAMES["special"], before_pt=6, after_pt=5, size_pt=11, bold=True)
    configure_style(doc, STYLE_NAMES["note"], left_cm=0.6, right_cm=0.6, after_pt=5, size_pt=10, italic=True)


def add_text_paragraph(doc, text, block_type):
    text = str(text or "").strip()
    paragraph = doc.add_paragraph()

    if block_type == "scene_heading":
        paragraph.style = STYLE_NAMES["scene_heading"]
        set_paragraph_format(paragraph, before_pt=10, after_pt=8, align=WD_ALIGN_PARAGRAPH.LEFT)
        run = paragraph.add_run(text)
        set_run_font(run, 12, bold=True)
    elif block_type == "action":
        paragraph.style = STYLE_NAMES["action"]
        set_paragraph_format(paragraph, after_pt=7, align=WD_ALIGN_PARAGRAPH.LEFT)
        run = paragraph.add_run(text)
        set_run_font(run, 11)
    elif block_type in {"character", "speaker"}:
        paragraph.style = STYLE_NAMES["speaker"]
        set_paragraph_format(paragraph, left_cm=6.6, after_pt=1, align=WD_ALIGN_PARAGRAPH.LEFT)
        run = paragraph.add_run(text)
        set_run_font(run, 11, bold=True)
    elif block_type == "dialogue":
        paragraph.style = STYLE_NAMES["dialogue"]
        set_paragraph_format(paragraph, left_cm=4.1, right_cm=3.6, after_pt=5, align=WD_ALIGN_PARAGRAPH.LEFT)
        run = paragraph.add_run(text)
        set_run_font(run, 11)
    elif block_type == "parenthetical":
        paragraph.style = STYLE_NAMES["parenthetical"]
        set_paragraph_format(paragraph, left_cm=5.0, right_cm=4.0, after_pt=3, align=WD_ALIGN_PARAGRAPH.LEFT)
        run = paragraph.add_run(text)
        set_run_font(run, 10)
    elif block_type == "transition":
        paragraph.style = STYLE_NAMES["transition"]
        set_paragraph_format(paragraph, after_pt=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
        run = paragraph.add_run(text)
        set_run_font(run, 11)
    elif block_type in {"shot", "insert", "super", "montage", "intercut"}:
        paragraph.style = STYLE_NAMES["special"]
        set_paragraph_format(paragraph, before_pt=6, after_pt=5, align=WD_ALIGN_PARAGRAPH.LEFT)
        run = paragraph.add_run(text)
        set_run_font(run, 11, bold=True)
    elif block_type == "note":
        paragraph.style = STYLE_NAMES["note"]
        set_paragraph_format(paragraph, left_cm=0.6, right_cm=0.6, after_pt=5, align=WD_ALIGN_PARAGRAPH.LEFT)
        run = paragraph.add_run(text)
        set_run_font(run, 10)
        run.font.italic = True
    else:
        paragraph.style = STYLE_NAMES["action"]
        set_paragraph_format(paragraph, after_pt=6, align=WD_ALIGN_PARAGRAPH.LEFT)
        run = paragraph.add_run(text)
        set_run_font(run, 11)


def add_page_number_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run._r.append(text)
    run._r.append(fld_char_end)
    set_run_font(run, 10)


def setup_document(title=None):
    doc = Document()
    if title:
        doc.core_properties.title = title
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    add_page_number_footer(section)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "SimSun"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    configure_screenplay_styles(doc)

    return doc


def load_character_name_map(film_data_dir, issues):
    story_bible_path = film_data_dir / "story-bible.yaml"
    if not story_bible_path.exists():
        issues.append(
            make_issue(
                "warning",
                "story_bible_missing",
                "story-bible.yaml not found; dialogue speakers may fall back to character_id.",
                "story-bible.yaml",
            )
        )
        return {}

    story_bible = read_yaml(story_bible_path, issues)
    if not isinstance(story_bible, dict):
        return {}

    name_map = {}
    for character in story_bible.get("characters") or []:
        if not isinstance(character, dict):
            continue
        character_id = character.get("character_id") or character.get("id")
        name = character.get("name") or character.get("display_name")
        if character_id and name:
            name_map[str(character_id)] = str(name)
    return name_map


def resolve_blocks_for_ref(film_data_dir, page_ref, issues):
    script_path = page_ref.get("script_path")
    start_id = page_ref.get("start_block_id")
    end_id = page_ref.get("end_block_id")
    if not script_path or not start_id or not end_id:
        issues.append(make_issue("error", "invalid_block_ref", "block_ref must include script_path, start_block_id, and end_block_id."))
        return []

    block_path = film_data_dir / script_path
    if not block_path.exists():
        issues.append(make_issue("error", "script_path_missing", "Referenced script_path does not exist.", script_path))
        return []

    blocks = read_yaml(block_path, issues)
    if not isinstance(blocks, list):
        issues.append(make_issue("error", "blocks_not_list", "Referenced blocks.yaml must be a list.", script_path))
        return []

    start_index = end_index = None
    for index, block in enumerate(blocks):
        if not isinstance(block, dict):
            continue
        block_id = block.get("block_id")
        if block_id == start_id:
            start_index = index
        if block_id == end_id:
            end_index = index

    if start_index is None or end_index is None:
        issues.append(make_issue("error", "block_range_missing", f"Cannot find block range {start_id} -> {end_id}.", script_path))
        return []
    if start_index > end_index:
        issues.append(make_issue("error", "block_range_reversed", f"Block range is reversed: {start_id} -> {end_id}.", script_path))
        return []

    return [(block_path, block) for block in blocks[start_index : end_index + 1] if isinstance(block, dict)]


def speaker_text(block, character_name_map):
    text = block.get("speaker") or block.get("character_name") or block.get("character")
    if text:
        return str(text)
    character_id = block.get("character_id")
    if character_id and str(character_id) in character_name_map:
        return character_name_map[str(character_id)]
    if character_id:
        return str(character_id).replace("char_", "").replace("_", " ").upper()
    return ""


def write_block(doc, block, character_name_map):
    block_type = str(block.get("block_type", "")).strip()
    text = str(block.get("text", "") or "")
    if block_type == "dialogue":
        speaker = speaker_text(block, character_name_map)
        if speaker:
            add_text_paragraph(doc, speaker, "speaker")
        add_text_paragraph(doc, text, "dialogue")
        return
    add_text_paragraph(doc, text, block_type)


def export_docx(
    film_data_dir,
    output,
    title=None,
    strict=False,
    render_check_dir=None,
    show_page_labels=False,
    page_break_per_script_page=False,
):
    issues = []
    film_yaml = film_data_dir / "film.yaml"
    pages_yaml = film_data_dir / "script" / "pages.yaml"

    film = read_yaml(film_yaml, issues) if film_yaml.exists() else None
    if not film_yaml.exists():
        issues.append(make_issue("warning", "film_yaml_missing", "film.yaml not found; title metadata unavailable.", "film.yaml"))
    if not pages_yaml.exists():
        issues.append(make_issue("error", "pages_yaml_missing", "Root script/pages.yaml is required for DOCX export.", "script/pages.yaml"))
        return {"ok": False, "output": str(output), "issues": issues, "totals": {}}

    pages = read_yaml(pages_yaml, issues)
    if not isinstance(pages, list):
        issues.append(make_issue("error", "pages_not_list", "script/pages.yaml must be a list.", "script/pages.yaml"))
        return {"ok": False, "output": str(output), "issues": issues, "totals": {}}

    if not title and isinstance(film, dict):
        title = film.get("title")

    doc = setup_document(title=title)
    character_name_map = load_character_name_map(film_data_dir, issues)
    seen_block_keys = {}
    scene_heading_cache = {}
    page_count = 0
    block_count = 0

    for page in sorted((p for p in pages if isinstance(p, dict)), key=lambda p: p.get("page_number", 0)):
        page_count += 1
        page_number = page.get("page_number", page_count)
        if show_page_labels:
            heading = doc.add_paragraph()
            set_paragraph_format(heading, after_pt=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
            run = heading.add_run(f"第 {page_number} 页")
            set_run_font(run, 9)
            run.font.italic = True

        refs = page.get("block_refs") or []
        if not isinstance(refs, list):
            issues.append(make_issue("error", "block_refs_not_list", f"block_refs must be a list on page {page_number}."))
            refs = []

        for page_ref in refs:
            resolved_blocks = resolve_blocks_for_ref(film_data_dir, page_ref, issues)
            if page_break_per_script_page and resolved_blocks:
                script_path = page_ref.get("script_path")
                start_id = page_ref.get("start_block_id")
                first_path, first_block = resolved_blocks[0]
                if first_block.get("block_type") != "scene_heading":
                    if first_path not in scene_heading_cache:
                        all_blocks = read_yaml(first_path, issues)
                        scene_heading_cache[first_path] = ""
                        if isinstance(all_blocks, list):
                            for candidate in all_blocks:
                                if isinstance(candidate, dict) and candidate.get("block_type") == "scene_heading":
                                    scene_heading_cache[first_path] = str(candidate.get("text") or "").strip()
                                    break
                    heading_text = scene_heading_cache.get(first_path)
                    if heading_text:
                        add_text_paragraph(doc, f"{heading_text}（续）", "scene_heading")
                    elif start_id:
                        issues.append(
                            make_issue(
                                "warning",
                                "continuation_heading_missing",
                                f"ScriptPage starts mid-scene at {script_path}#{start_id}, but no scene heading was found for continuation.",
                                script_path,
                            )
                        )

            for block_path, block in resolved_blocks:
                block_id = block.get("block_id")
                key = f"{block_path.relative_to(film_data_dir).as_posix()}#{block_id}"
                if block_id:
                    if key in seen_block_keys:
                        issues.append(
                            make_issue(
                                "warning",
                                "duplicate_block_export",
                                f"Block {key} is referenced by multiple ScriptPages; clean root script/pages.yaml by splitting multi-page scenes into non-overlapping block ranges before formal DOCX export.",
                            )
                        )
                    else:
                        seen_block_keys[key] = page_number
                write_block(doc, block, character_name_map)
                block_count += 1

        if page_break_per_script_page and page_count < len(pages):
            doc.add_page_break()

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)

    if render_check_dir:
        issues.extend(render_docx(output, render_check_dir))

    error_count = sum(1 for issue in issues if issue["level"] == "error")
    if strict:
        duplicate_warnings = [issue for issue in issues if issue["type"] == "duplicate_block_export"]
        error_count += len(duplicate_warnings)

    return {
        "ok": error_count == 0,
        "output": str(output),
        "totals": {
            "scriptPages": len(pages),
            "exportedPages": page_count,
            "exportedBlocks": block_count,
            "issues": len(issues),
        },
        "issues": issues,
    }


def render_docx(output, render_check_dir):
    issues = []
    render_check_dir.mkdir(parents=True, exist_ok=True)
    soffice = shutil.which("soffice")
    pdftoppm = shutil.which("pdftoppm")
    if not soffice:
        issues.append(make_issue("warning", "soffice_missing", "LibreOffice soffice not found; skipped visual render check."))
        return issues
    if not pdftoppm:
        issues.append(make_issue("warning", "pdftoppm_missing", "pdftoppm not found; skipped PNG render check."))
        return issues

    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(render_check_dir), str(output)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        pdf_path = render_check_dir / f"{output.stem}.pdf"
        if not pdf_path.exists():
            issues.append(make_issue("warning", "pdf_render_missing", "LibreOffice did not produce expected PDF."))
            return issues
        subprocess.run(
            [pdftoppm, "-png", "-f", "1", "-l", "3", str(pdf_path), str(render_check_dir / output.stem)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
    except Exception as exc:
        issues.append(make_issue("warning", "render_check_failed", f"Visual render check failed: {exc}"))
    return issues


def print_human(report):
    print(f"Output: {report['output']}")
    print(f"OK: {'yes' if report['ok'] else 'no'}")
    for key, value in report.get("totals", {}).items():
        print(f"{key}: {value}")
    if report["issues"]:
        print("Issues:")
        for issue in report["issues"]:
            print(f"- [{issue['level']}] {issue['type']}: {issue['message']}")


def main(argv=None):
    parser = argparse.ArgumentParser(description="Export local film screenplay pages to DOCX.")
    parser.add_argument("--film-data-dir", required=True, help="Path to local film-data directory")
    parser.add_argument("--output", required=True, help="Output .docx path")
    parser.add_argument("--title", help="Optional document title paragraph")
    parser.add_argument("--strict", action="store_true", help="Return non-zero for errors and duplicate block exports")
    parser.add_argument("--render-check-dir", help="Optional directory for DOCX->PDF->PNG render check")
    parser.add_argument("--show-page-labels", action="store_true", help="Print ScriptPage labels in body for debugging")
    parser.add_argument(
        "--page-break-per-script-page",
        action="store_true",
        help="Debug/audit mode: force a physical page break after each ScriptPage and add continuation headings when a page starts mid-scene",
    )
    parser.add_argument("--json", action="store_true", help="Print JSON report")
    args = parser.parse_args(argv)

    if not ensure_docx_available():
        return 2

    report = export_docx(
        Path(args.film_data_dir).resolve(),
        Path(args.output).resolve(),
        title=args.title,
        strict=args.strict,
        render_check_dir=Path(args.render_check_dir).resolve() if args.render_check_dir else None,
        show_page_labels=args.show_page_labels,
        page_break_per_script_page=args.page_break_per_script_page,
    )

    if args.json:
        print(json.dumps(report, ensure_ascii=False, indent=2))
    else:
        print_human(report)
    return 0 if report["ok"] else 1


if __name__ == "__main__":
    sys.exit(main())
